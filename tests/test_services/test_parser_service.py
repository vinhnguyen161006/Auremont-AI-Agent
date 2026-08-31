import fitz
import pytest
from docx import Document

from backend.services.parser_service import (
    UnsupportedDocumentTypeError,
    parse_document,
)


def _make_pdf_bytes() -> bytes:
    document = fitz.open()

    first_page = document.new_page()
    first_page.insert_text((72, 72), "Bang gia can ho 2PN")

    second_page = document.new_page()
    second_page.insert_text((72, 72), "Chinh sach thanh toan")

    data = document.tobytes()
    document.close()
    return data


def test_parse_pdf_keeps_page_numbers():
    sections = parse_document("bang-gia.pdf", _make_pdf_bytes())

    assert len(sections) == 2
    assert sections[0].page == 1
    assert "Bang gia" in sections[0].text
    assert sections[1].page == 2


def _make_pdf_with_ruled_table_bytes() -> bytes:
    """A page with prose, then a ruled 2-column label/value table, then more prose —
    mirrors the label/value tables common in Vietnamese real-estate brochures (e.g.
    "Ten du an | The Senique Hanoi")."""
    document = fitz.open()
    page = document.new_page()

    rows = [
        ("Ten du an", "The Senique Hanoi"),
        ("Chu dau tu", "CapitaLand Development"),
        ("Vi tri", "O B6-CT02 Gia Lam"),
    ]

    page.insert_text((50, 50), "Day la doan van truoc bang.", fontsize=11)

    y = 80
    row_height = 25
    table_top = 65
    table_bottom = table_top + len(rows) * row_height + 5
    for label, value in rows:
        page.insert_text((50, y), label, fontsize=11)
        page.insert_text((250, y), value, fontsize=11)
        y += row_height

    page.draw_rect(fitz.Rect(40, table_top, 480, table_bottom))
    for index in range(len(rows) + 1):
        line_y = table_top + index * row_height + (5 if index == len(rows) else 0)
        page.draw_line((40, line_y), (480, line_y))
    page.draw_line((240, table_top), (240, table_bottom))

    page.insert_text((50, y + 20), "Day la doan van sau bang.", fontsize=11)

    data = document.tobytes()
    document.close()
    return data


def test_parse_pdf_extracts_ruled_table_as_markdown():
    sections = parse_document("thong-tin-du-an.pdf", _make_pdf_with_ruled_table_bytes())

    table_sections = [section for section in sections if section.content_type == "table"]
    assert len(table_sections) == 1

    table_text = table_sections[0].text
    assert table_sections[0].page == 1
    assert "Ten du an" in table_text
    assert "The Senique Hanoi" in table_text
    assert "Chu dau tu" in table_text
    assert "CapitaLand Development" in table_text
    assert "Vi tri" in table_text
    assert "O B6-CT02 Gia Lam" in table_text


def test_parse_pdf_excludes_table_text_from_prose_section():
    sections = parse_document("thong-tin-du-an.pdf", _make_pdf_with_ruled_table_bytes())

    prose_sections = [section for section in sections if section.content_type == "prose"]
    assert len(prose_sections) == 1

    prose_text = prose_sections[0].text
    assert "Day la doan van truoc bang." in prose_text
    assert "Day la doan van sau bang." in prose_text
    assert "The Senique Hanoi" not in prose_text
    assert "CapitaLand Development" not in prose_text


def test_parse_pdf_without_tables_is_unchanged():
    sections = parse_document("bang-gia.pdf", _make_pdf_bytes())

    assert len(sections) == 2
    assert all(section.content_type == "prose" for section in sections)


def test_parse_docx_returns_text_without_page(tmp_path):
    path = tmp_path / "chinh-sach.docx"

    document = Document()
    document.add_paragraph("Chinh sach ban hang")
    document.add_paragraph("Ho tro vay ngan hang")
    document.save(path)

    sections = parse_document(path.name, path.read_bytes())

    assert len(sections) == 1
    assert sections[0].page is None
    assert "Chinh sach ban hang" in sections[0].text


def test_parse_docx_reads_tables_in_document_order(tmp_path):
    """Bảng chiết khấu/tiến độ trong file Word phải vào được kho tri thức.

    `document.paragraphs` bỏ qua bảng, nên trước đây file ingest "thành công" mà
    toàn bộ con số biến mất — không có cảnh báo nào cho Admin.
    """
    path = tmp_path / "csbh.docx"

    document = Document()
    document.add_paragraph("II. CHINH SACH CHIET KHAU")
    table = document.add_table(rows=3, cols=3)
    for row_index, cells in enumerate(
        [
            ("Loai can", "Chiet khau", "Dieu kien"),
            ("2PN", "5%", "Thanh toan som 95%"),
            ("3PN", "7%", "Thanh toan som 95%"),
        ]
    ):
        for column_index, value in enumerate(cells):
            table.cell(row_index, column_index).text = value
    document.add_paragraph("Uu dai khong cong don voi chuong trinh khac.")
    document.save(path)

    sections = parse_document(path.name, path.read_bytes())
    text = sections[0].text

    assert "2PN" in text and "5%" in text
    assert "3PN" in text and "7%" in text
    assert text.index("CHINH SACH CHIET KHAU") < text.index("2PN") < text.index("khong cong don")
    assert "| Loai can | Chiet khau | Dieu kien |" in text
    assert "| --- | --- | --- |" in text


def test_parse_docx_with_only_a_table_is_not_treated_as_empty(tmp_path):
    """Bảng giá thuần (không có đoạn văn nào) trước đây bị coi là file rỗng."""
    path = tmp_path / "bang-gia.docx"

    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Ma can"
    table.cell(0, 1).text = "Gia"
    table.cell(1, 0).text = "BE1-1201"
    table.cell(1, 1).text = "3.5 ty"
    document.save(path)

    sections = parse_document(path.name, path.read_bytes())

    assert "BE1-1201" in sections[0].text
    assert "3.5 ty" in sections[0].text


def test_parse_rejects_unsupported_extension():
    with pytest.raises(UnsupportedDocumentTypeError):
        parse_document("bang-gia.xlsx", b"not relevant")
