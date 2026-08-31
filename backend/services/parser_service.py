import logging
from collections.abc import Iterator
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedSection:
    """A block of text extracted from a source document.

    PDFs keep their page number (1-based); DOCX has no stable notion of a page,
    so None is used there.
    """

    text: str
    page: int | None
    content_type: str = "prose"
    block_offsets: tuple[tuple[int, float], ...] = ()


class UnsupportedDocumentTypeError(ValueError):
    """The file is not in a format the pipeline supports."""


class DocumentParseError(ValueError):
    """The file is corrupt or no text could be extracted from it."""


def parse_document(filename: str, data: bytes) -> list[ParsedSection]:
    """Parse PDF/DOCX straight from bytes, with no dependency on a temp file on disk."""
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(data)

    if suffix == ".docx":
        return _parse_docx(data)

    raise UnsupportedDocumentTypeError(
        f"Unsupported file type: {suffix or '(no extension)'}. Only .pdf and .docx are supported."
    )


def _parse_pdf(data: bytes) -> list[ParsedSection]:
    try:
        with fitz.open(stream=data, filetype="pdf") as pdf:
            sections: list[ParsedSection] = []
            for page_index, page in enumerate(pdf):
                page_no = page_index + 1
                table_sections, table_bboxes = _extract_tables(page, page_no)
                prose_text = _extract_prose_excluding(page, table_bboxes).strip()

                if prose_text:
                    sections.append(
                        ParsedSection(
                            text=prose_text,
                            page=page_no,
                            block_offsets=_block_offsets(page, prose_text),
                        )
                    )
                sections.extend(table_sections)
    except Exception as exc:
        logger.exception(
            "PDF parsing failed.",
            extra={"event": "parser.pdf.failed", "size_bytes": len(data)},
        )
        raise DocumentParseError("Could not parse PDF.") from exc

    if not sections:
        raise DocumentParseError("PDF contains no extractable text. OCR is not supported yet.")

    return sections


def _extract_tables(page: "fitz.Page", page_no: int) -> tuple[list[ParsedSection], list["fitz.Rect"]]:
    """Detect ruled tables on a page and render each as a markdown ParsedSection.

    A find_tables() failure on one page degrades to "no tables found here" rather
    than aborting the whole document — table detection is a best-effort enhancement,
    not something the rest of parsing should depend on succeeding.
    """
    try:
        finder = page.find_tables()
    except Exception:
        logger.warning(
            "Table detection failed on page %s; treating page as prose.",
            page_no,
            exc_info=True,
            extra={"event": "parser.pdf.table_detection_failed", "page": page_no},
        )
        return [], []

    sections: list[ParsedSection] = []
    bboxes: list[fitz.Rect] = []

    for table in finder.tables:
        if table.row_count < 2 or table.col_count < 2:
            continue

        markdown = _rows_to_markdown(table.extract())
        if not markdown:
            continue

        sections.append(ParsedSection(text=markdown, page=page_no, content_type="table"))
        bboxes.append(fitz.Rect(table.bbox))

    return sections, bboxes


def _rows_to_markdown(rows: list[list[str | None]]) -> str:
    """Render raw extracted rows as a markdown table without promoting any row to
    a header. Many tables in Vietnamese real-estate brochures are label/value pairs
    with no real header row (e.g. "Tên dự án | The Senique Hanoi"); PyMuPDF's own
    to_markdown() would silently treat the first data row as a header and lose it
    from the row data, so the markdown is built manually here instead.
    """
    cleaned = [[(cell or "").strip().replace("\n", " ") for cell in row] for row in rows]
    cleaned = [row for row in cleaned if any(row)]
    if not cleaned:
        return ""

    col_count = max(len(row) for row in cleaned)
    lines = ["|" + "|".join(row + [""] * (col_count - len(row))) + "|" for row in cleaned]
    return "\n".join(lines)


def _extract_prose_excluding(page: "fitz.Page", table_bboxes: list["fitz.Rect"]) -> str:
    """Reconstruct page text with any block overlapping a table region removed.

    With no detected tables this reproduces page.get_text("text") exactly, so pages
    without tables are provably unchanged by this feature.
    """
    if not table_bboxes:
        return page.get_text("text")

    lines_out: list[str] = []
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:
            continue

        block_rect = fitz.Rect(block["bbox"])
        if _overlaps_any_table(block_rect, table_bboxes):
            continue

        for line in block["lines"]:
            line_text = "".join(span["text"] for span in line["spans"])
            if line_text.strip():
                lines_out.append(line_text)

    return "\n".join(lines_out)


def _overlaps_any_table(block_rect: "fitz.Rect", table_bboxes: list["fitz.Rect"], threshold: float = 0.5) -> bool:
    area = block_rect.get_area()
    if area <= 0:
        return False
    return any((block_rect & bbox).get_area() / area > threshold for bbox in table_bboxes)


def _block_offsets(page: "fitz.Page", page_text: str) -> tuple[tuple[int, float], ...]:
    """Best-effort (offset into `page_text`, Y-from-top in PDF points) breakpoints.

    `page.get_text("blocks")` gives each visual block's bounding box; `y0` (top edge,
    increasing downward — verified against this project's own PDFs, not assumed) is what
    a citation later scrolls a viewer to. Matched onto `page_text` by searching for each
    block's own first line: blocks mode and text mode don't promise identical whitespace,
    so a block whose text can't be found is just skipped — that block's chunk falls back
    to page-top on citation, exactly today's behaviour, rather than raising.
    """
    breakpoints: list[tuple[int, float]] = []
    for block in page.get_text("blocks"):
        x0, y0, x1, y1, block_text, block_no, block_type = block
        if block_type != 0:
            continue
        first_line = block_text.strip().splitlines()[0].strip() if block_text.strip() else ""
        if not first_line:
            continue
        offset = page_text.find(first_line)
        if offset >= 0:
            breakpoints.append((offset, y0))
    breakpoints.sort(key=lambda item: item[0])
    return tuple(breakpoints)


def _parse_docx(data: bytes) -> list[ParsedSection]:
    """Extract prose AND tables, in the order they appear in the document.

    `document.paragraphs` skips tables entirely. For this corpus that is not a detail:
    discount tiers, payment schedules and price rows in a .docx CSBH live in tables and
    nowhere else, so reading only paragraphs ingests the file "successfully" while every
    figure in it silently disappears — the worst possible failure, because nothing looks
    wrong until a Sale is told the policy has no discount data.
    """
    try:
        document = DocxDocument(BytesIO(data))
        blocks = []
        for item in _iter_block_items(document):
            rendered = _render_table(item) if isinstance(item, Table) else " ".join(item.text.split())
            if rendered:
                blocks.append(rendered)
        text = "\n\n".join(blocks)
    except Exception as exc:
        logger.exception(
            "DOCX parsing failed.",
            extra={"event": "parser.docx.failed", "size_bytes": len(data)},
        )
        raise DocumentParseError("Could not parse DOCX.") from exc

    if not text:
        raise DocumentParseError("DOCX contains no extractable text.")

    return [ParsedSection(text=text, page=None)]


def _iter_block_items(document) -> Iterator[Paragraph | Table]:
    """Walk the document body in reading order.

    python-docx exposes `.paragraphs` and `.tables` as two separate flat lists, which
    loses the interleaving. A discount table has to stay attached to the clause that
    introduces it, otherwise the chunk carrying the numbers has no idea what they apply to.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _render_table(table: Table) -> str:
    """Render a table as pipe-delimited rows with a header separator.

    This is the shape `chunking_service` already recognises (`TABLE_ROW_RE`), so a Word
    table lands in the same table-aware path as a Markdown one: split only between rows,
    with the column header repeated in every chunk.
    """
    rows: list[str] = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        if not any(cells):
            continue
        rows.append("| " + " | ".join(cells) + " |")

    if not rows:
        return ""

    column_count = rows[0].count("|") - 1
    separator = "| " + " | ".join(["---"] * column_count) + " |"
    return "\n".join([rows[0], separator, *rows[1:]])
